import json
import os
from docx import Document

def extract_full_text(doc):
    content = []

    # Extract normal paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)

    # Extract text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    return content


def convert_docx_to_json():
    print("Starting conversion...")

    docx_path = "/homes/neumann/teklehaymanot/TigrinyaTokenizer/MPETokenization/Paralleldata/MoVoC/movoc/Untitled spreadsheet - Sheet1.pdf"

    # Load document
    document = Document(docx_path)

    # Extract ALL text
    content = extract_full_text(document)

    # Prepare JSON
    data = {
        "filename": os.path.basename(docx_path),
        "content": content
    }

    # Save JSON next to original file
    json_path = os.path.splitext(docx_path)[0] + ".json"

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

    print(f"JSON file saved to: {json_path}")


if __name__ == "__main__":
    convert_docx_to_json()
